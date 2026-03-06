# generate_final_srs_ai_resume_screener.py
# Outputs:
#   - SRS_AI_Resume_Screener_Chaynika_Vyas_Final.docx
#   - SRS_AI_Resume_Screener_Chaynika_Vyas_Final.pdf  (if Word installed; docx2pdf)

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from graphviz import Digraph
from docx2pdf import convert as docx2pdf_convert
import os
import shutil
# -------------------------------------------------------------------------
# Auto-detect Graphviz and fix PATH if 'dot' not found
# -------------------------------------------------------------------------
import os, shutil, sys

def ensure_graphviz_path():
    """Ensure Graphviz executables (dot, neato, etc.) are accessible."""
    dot_path = shutil.which("dot")
    if dot_path:
        print(f"[INFO] Graphviz found: {dot_path}")
        return True

    # Common Graphviz installation paths
    possible_paths = [
        r"C:\Program Files\Graphviz\bin",
        r"C:\Program Files (x86)\Graphviz\bin",
        r"C:\Users\\" + os.getlogin() + r"\AppData\Local\Programs\Graphviz\bin"
    ]

    found = False
    for path in possible_paths:
        if os.path.exists(os.path.join(path, "dot.exe")):
            os.environ["PATH"] += os.pathsep + path
            print(f"[INFO] Added Graphviz to PATH: {path}")
            found = True
            break

    if not found:
        print("[WARNING] Graphviz not found in common paths.")
        print("Please ensure Graphviz is installed and its 'bin' folder is accessible.")
        return False
    return True

# Call this before using graphviz
ensure_graphviz_path()


AUTHOR = "Chaynika Vyas"
COLLEGE = "SGSITS, Indore"
DEPARTMENT = "Computer Technology and Application"
HEADER_TEXT = "AI Resume Screener – SRS Report"
DOCX_OUT = "SRS_AI_Resume_Screener_Chaynika_Vyas_Final.docx"
PDF_OUT  = "SRS_AI_Resume_Screener_Chaynika_Vyas_Final.pdf"

# Put your 6 screenshots here (exact filenames in your folder):
SCREENSHOTS = [
    "Screenshot 2025-11-10 151212.png",
    "Screenshot 2025-11-10 151249.png",
    "Screenshot 2025-11-10 151257.png",
    "Screenshot 2025-11-10 151310.png",
    "Screenshot 2025-11-10 151321.png",
    "Screenshot 2025-11-10 151330.png",
]

def set_doc_defaults(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

def add_header_footer(doc: Document, header_text: str):
    for section in doc.sections:
        # Header (center)
        header = section.header.paragraphs[0]
        header.text = header_text
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Footer page number (bottom-right)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run()
        fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = " PAGE "
        fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)

def H(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0,0,0)
    return p

def P(doc, text):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return para

def bullet(doc, items):
    for it in items:
        para = doc.add_paragraph(it)
        para.style = doc.styles['List Paragraph']
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_image(doc, path, width_in=5.8):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_in))
    else:
        P(doc, f"(Image missing: {path})")

# ---------- Graphviz helpers ----------
def ensure_graphviz_available():
    # Quick check for dot.exe on Windows PATH
    return shutil.which("dot") is not None

def save_graphviz_png(graph: Digraph, filename: str):
    graph.format = "png"
    graph.render(filename, cleanup=True)  # creates filename.png

def diagram_system_architecture(out="diagram_architecture"):
    g = Digraph("Architecture", graph_attr={'rankdir':'LR','fontsize':'10'})
    g.node("U","User (Recruiter)\nBrowser", shape="box", style="rounded,filled", fillcolor="#E3F2FD")
    g.node("S","Streamlit Frontend", shape="box", style="rounded,filled", fillcolor="#E8F5E9")
    g.node("P","Parser Layer\n(pdfminer / python-docx)", shape="box", style="rounded,filled", fillcolor="#FFF3E0")
    g.node("N","NLP + Scoring\n(NLTK / TF-IDF / Cosine)", shape="box", style="rounded,filled", fillcolor="#F3E5F5")
    g.node("D","Data Store\n(CSV/SQLite)", shape="cylinder", style="filled", fillcolor="#E1F5FE")
    g.edges([("U","S"),("S","P"),("P","N"),("N","S"),("N","D")])
    save_graphviz_png(g, out); return out + ".png"

def diagram_dfd_l0(out="diagram_dfd_l0"):
    g = Digraph("DFD_L0", graph_attr={'rankdir':'LR','fontsize':'10'})
    g.node("R","Recruiter", shape="oval", style="filled", fillcolor="#FFFDE7")
    g.node("SYS","AI Resume Screener", shape="box", style="rounded,filled", fillcolor="#E8EAF6")
    g.node("JD","Job Description", shape="box", style="rounded,filled", fillcolor="#E1F5FE")
    g.node("RES","Resumes", shape="box", style="rounded,filled", fillcolor="#E1F5FE")
    g.node("OUT","Ranked Results", shape="box", style="rounded,filled", fillcolor="#E1F5FE")
    g.edges([("R","SYS"),("JD","SYS"),("RES","SYS"),("SYS","OUT"),("SYS","R")])
    save_graphviz_png(g, out); return out + ".png"

def diagram_erd(out="diagram_erd"):
    g = Digraph("ERD", graph_attr={'rankdir':'LR','fontsize':'10'})
    g.node("JD","job_descriptions| jd_id (PK)\\l title\\l jd_text\\l", shape="record", style="filled", fillcolor="#E3F2FD")
    g.node("RS","resumes| resume_id (PK)\\l name\\l file_path\\l parsed_text\\l", shape="record", style="filled", fillcolor="#E8F5E9")
    g.node("RE","results| result_id (PK)\\l resume_id (FK)\\l jd_id (FK)\\l score\\l missing_keywords\\l created_at\\l", shape="record", style="filled", fillcolor="#FFF3E0")
    g.edge("JD","RE", label="1..*")
    g.edge("RS","RE", label="1..*")
    save_graphviz_png(g, out); return out + ".png"

def diagram_usecase(out="diagram_usecase"):
    g = Digraph("UseCase", graph_attr={'rankdir':'LR','fontsize':'10'})
    g.node("Actor","Recruiter", shape="oval", style="filled", fillcolor="#FFFDE7")
    g.node("UC","AI Resume Screener", shape="box", style="rounded,filled", fillcolor="#E8EAF6")
    g.edge("Actor","UC", label="Upload JD & Resumes / View Results")
    save_graphviz_png(g, out); return out + ".png"

def diagram_class(out="diagram_class"):
    g = Digraph("ClassDiagram", graph_attr={'rankdir':'TB','fontsize':'10'})
    g.node("Parser","Parser|+ extract_text(file)\\l+ parse_pdf()\\l+ parse_docx()\\l", shape="record", style="filled", fillcolor="#E3F2FD")
    g.node("NLP","NLPService|+ preprocess(text)\\l+ tfidf_score()\\l", shape="record", style="filled", fillcolor="#E8F5E9")
    g.node("Repo","ResultRepo|+ save(result)\\l+ to_csv()\\l", shape="record", style="filled", fillcolor="#FFF3E0")
    g.node("App","App|+ run()\\l", shape="record", style="filled", fillcolor="#F3E5F5")
    g.edges([("App","Parser"),("App","NLP"),("NLP","Repo")])
    save_graphviz_png(g, out); return out + ".png"

def diagram_state(out="diagram_state"):
    g = Digraph("State", graph_attr={'rankdir':'LR','fontsize':'10'})
    g.node("Idle","Idle", shape="circle", style="filled", fillcolor="#E0F2F1")
    g.node("Up","Uploading Files", shape="box", style="rounded,filled", fillcolor="#E8F5E9")
    g.node("Proc","Processing NLP", shape="box", style="rounded,filled", fillcolor="#E3F2FD")
    g.node("Show","Showing Results", shape="box", style="rounded,filled", fillcolor="#FFF3E0")
    g.edges([("Idle","Up"),("Up","Proc"),("Proc","Show"),("Show","Idle")])
    save_graphviz_png(g, out); return out + ".png"

def diagram_activity(out="diagram_activity"):
    g = Digraph("Activity", graph_attr={'rankdir':'TB','fontsize':'10'})
    g.node("Start","Start", shape="circle")
    g.node("U","Upload JD & Resumes", shape="box", style="rounded,filled", fillcolor="#E8EAF6")
    g.node("E","Extract Text", shape="box", style="rounded,filled", fillcolor="#E3F2FD")
    g.node("C","TF-IDF + Cosine", shape="box", style="rounded,filled", fillcolor="#E8F5E9")
    g.node("R","Rank + CSV Report", shape="box", style="rounded,filled", fillcolor="#FFF3E0")
    g.node("End","End", shape="doublecircle")
    g.edges([("Start","U"),("U","E"),("E","C"),("C","R"),("R","End")])
    save_graphviz_png(g, out); return out + ".png"

# ---------- Content sections ----------
def H1(doc, t): return H(doc, t, 1)
def H2(doc, t): return H(doc, t, 2)

def section_1(doc):
    H1(doc, "1. Introduction")
    P(doc, "1.1 Purpose\nThis SRS defines complete requirements for the AI Resume Screener.")
    P(doc, "1.2 Scope\nAutomates comparison of resumes with a Job Description using NLP and similarity scoring, outputting rankings and keyword insights.")
    P(doc, "1.3 Problem Statement\nManual screening is slow, inconsistent, and subjective. A data-driven solution improves speed and consistency.")
    P(doc, "1.4 Objectives\n• Parse resumes (PDF/DOCX/TXT)\n• Extract & preprocess text\n• Compute similarity & rank\n• Highlight missing keywords\n• Export results as CSV")
    P(doc, "1.5 Definitions, Acronyms & Abbreviations\nNLP, TF-IDF, JD, HR")
    P(doc, "1.6 References\nPython/Streamlit/scikit-learn/NLTK docs; IEEE 830 SRS guidelines; IR/NLP literature.")

def section_2(doc):
    H1(doc, "2. Overall Description")
    P(doc, "2.1 Product Perspective\nStandalone Streamlit app; frontend for uploads/visualization; backend modules for parsing and NLP scoring.")
    P(doc, "2.2 Product Features\n• Resume/JD upload • Text extraction • NLP preprocessing • TF-IDF+cosine • Ranking+CSV • Keyword highlight.")
    P(doc, "2.3 User Classes and Characteristics\nPrimary: Recruiters/HR; Secondary: candidates/analysts; basic computer skills assumed.")
    P(doc, "2.4 Operating Environment\nWindows/macOS/Linux; Python 3.13; Streamlit; Chrome/Edge/Firefox.")
    P(doc, "2.5 Design Constraints\nText-based PDFs only (OCR needed for scans); memory/CPU usage grows with batch size.")
    P(doc, "2.6 Assumptions and Dependencies\nReadable files, accurate JD; availability of required Python libraries.")

def section_3(doc):
    H1(doc, "3. System Design")
    H2(doc, "3.1 System Architecture"); add_image(doc, diagram_system_architecture())
    H2(doc, "3.2 Data Flow Diagram — Level 0"); add_image(doc, diagram_dfd_l0())
    H2(doc, "3.3 Entity-Relationship Diagram (ERD)"); add_image(doc, diagram_erd())
    H2(doc, "3.4 Use Case Diagram"); add_image(doc, diagram_usecase())
    H2(doc, "3.5 Class Diagram"); add_image(doc, diagram_class())
    H2(doc, "3.6 State Diagram"); add_image(doc, diagram_state())
    H2(doc, "3.7 Activity Diagram"); add_image(doc, diagram_activity())

def section_4(doc):
    H1(doc, "4. Specific Requirements")
    H2(doc, "4.1 Functional Requirements")
    bullet(doc, [
        "User Registration (optional for multi-user deployments).",
        "Login/Logout (optional; for local demo may be skipped).",
        "Upload JD and multiple resumes.",
        "Parse & preprocess text.",
        "Compute TF-IDF vectors and cosine similarity.",
        "Rank candidates and show missing keywords.",
        "Download CSV report.",
    ])
    H2(doc, "4.2 Non-Functional Requirements")
    P(doc, "4.2.1 Performance\nSmall batches (5–20 resumes) should process within seconds on a typical laptop.")
    P(doc, "4.2.2 Security\nLocal processing by default; no third-party uploads unless enabled.")
    P(doc, "4.2.3 Usability\nLight theme, high-contrast headings, clear flows.")
    P(doc, "4.2.4 Reliability\nGraceful handling of unreadable files; consistent scoring.")
    P(doc, "4.2.5 Maintainability\nModular parser/NLP/app layers.")
    P(doc, "4.2.6 Scalability\nExtend with database, queues, or embeddings for larger sets.")
    P(doc, "4.2.7 Portability\nCross-platform via Python and browsers.")

def section_5(doc):
    H1(doc, "5. Interface and UI Design")
    P(doc, "5.1 User Interface Overview\nSidebar inputs (JD+resumes) and main results panel with download and preview.")
    P(doc, "5.2 UI Layout and Navigation Flow\nUpload → Run Screening → View ranked results → Preview highlights → Download CSV.")
    P(doc, "5.3 Color Scheme and Typography\nLight backgrounds with dark text; readable monospace in previews.")
    P(doc, "5.4 Responsive Design Strategy\nStreamlit wide layout; widgets adapt to window size.")
    P(doc, "5.5 Accessibility Considerations\nHigh-contrast headings, clear labeling, keyboard-friendly controls.")
    for path in SCREENSHOTS[:3]:
        add_image(doc, path)

def section_6(doc):
    H1(doc, "6. Use Case Descriptions")
    usecases = [
        ("Use Case 1 – User Registration", "Actor: Recruiter (optional)\nPre: System running\nMain: Enter details → create account\nPost: Account created."),
        ("Use Case 2 – User Login", "Actor: Recruiter\nMain: Enter credentials → access dashboard\nPost: Session active."),
        ("Use Case 3 – Create Post (Upload Set)", "Main: Provide JD + upload resumes\nPost: Files stored; ready for screening."),
        ("Use Case 4 – Delete Post (Remove Uploads)", "Main: Remove a specific uploaded resume/JD from current session."),
        ("Use Case 5 – Logout", "End session and clear temporary data."),
        ("Use Case 6 – View All Posts (Results)", "View ranked results, missing keywords, and preview highlights."),
    ]
    for title, desc in usecases:
        H2(doc, title); P(doc, desc)

def section_7(doc):
    H1(doc, "7. Testing and Validation")
    P(doc, "7.1 Testing Strategy\nUnit testing for parsing and preprocessing; integration tests for end-to-end flow; UAT with HR students.")
    P(doc, "7.2 Test Environment\nWindows 10/11, Python 3.13, Streamlit.")
    P(doc, "7.3 Unit Testing\nFunctions for extract_text, preprocess, and scoring validated.")
    P(doc, "7.4 Integration Testing\nUpload→Score→Report flow verified with multiple formats.")
    P(doc, "7.5 System Testing\nEdge cases: empty JD, corrupt PDFs, empty resumes.")
    P(doc, "7.6 User Acceptance Testing (UAT)\nEvaluators confirm clarity and relevance of ranking.")
    H2(doc, "7.7 Test Case Summary")
    bullet(doc, [
        "TC-01 Upload JD → JD stored → Pass",
        "TC-02 Upload Resume (PDF/DOCX/TXT) → Text parsed → Pass",
        "TC-03 Run Scoring → Score visible, ranking OK → Pass",
        "TC-04 CSV Export → File downloaded with rows & fields → Pass",
        "TC-05 Highlight Keywords → JD terms shown in preview → Pass"
    ])

def section_8(doc):
    H1(doc, "8. Hardware and Software Requirements")
    P(doc, "8.1 Hardware Requirements\n4–8 GB RAM recommended; multi-core CPU.")
    P(doc, "8.2 Software Requirements\nPython 3.13, Streamlit, scikit-learn, NLTK, pdfminer.six, python-docx, Graphviz.")
    P(doc, "8.3 Deployment Requirements\nLocal Streamlit run; optional cloud/container deployment.")

def section_9(doc):
    H1(doc, "9. Security and Risk Management")
    P(doc, "9.1 Security Mechanisms\nLocal-only processing; optional data cleanup post-session.")
    P(doc, "9.2 Risk Identification\nUnreadable scans; biased JD; large batch latency.")
    P(doc, "9.3 Risk Analysis and Mitigation\nAdd OCR; better JD guidance; limit batch size or add progress bars.")
    P(doc, "9.4 Data Protection and Privacy\nNo PII storage beyond session; follow institutional policy.")

def section_10(doc):
    H1(doc, "10. Future Scope")
    P(doc, "10.1 Planned Enhancements\nEmbeddings (SBERT), OCR, richer reports, recruiter notes.")
    P(doc, "10.2 Scalability Plans\nDatabase persistence (SQLite/PostgreSQL), batching, caching.")
    P(doc, "10.3 Integration with AI Features\nLLM summaries; spaCy skill extraction; cloud deploy.")

def section_11(doc):
    H1(doc, "11. Conclusion")
    P(doc, "11.1 Summary of Objectives\nAn efficient, objective AI-based screening tool defined by this SRS.")
    P(doc, "11.2 Final Remarks\nReady for academic evaluation and portfolio demonstration.")

def section_12(doc):
    H1(doc, "12. Appendices")
    H2(doc, "A. Tools and Technologies Used")
    bullet(doc, [
        "Python 3.13, Streamlit (UI), NLTK (tokenization/stopwords), scikit-learn (TF-IDF, cosine), pdfminer.six & python-docx (parsing), Graphviz (diagrams)."
    ])
    H2(doc, "B. Sample Screenshots / UI Mock-ups")
    for path in SCREENSHOTS[3:]:
        add_image(doc, path)
    H2(doc, "C. References and Bibliography")
    bullet(doc, [
        "Official docs: Python, Streamlit, scikit-learn, NLTK, pdfminer.six, python-docx, Graphviz.",
        "Text similarity & SRS standards (IEEE 830)."
    ])

def build_doc():
    # Check Graphviz availability (nice-to-have)
    if not ensure_graphviz_available():
        print("⚠ Graphviz 'dot' not found on PATH. Diagrams may fail. Install Graphviz and reopen terminal.")

    doc = Document()
    set_doc_defaults(doc)
    add_header_footer(doc, HEADER_TEXT)

    # Sections 1–12
    section_1(doc)
    section_2(doc)
    section_3(doc)
    section_4(doc)
    section_5(doc)
    section_6(doc)
    section_7(doc)
    section_8(doc)
    section_9(doc)
    section_10(doc)
    section_11(doc)
    section_12(doc)

    doc.save(DOCX_OUT)
    print(f"✅ DOCX generated: {DOCX_OUT}")

    # Try PDF export (requires MS Word installed)
    try:
        docx2pdf_convert(DOCX_OUT, PDF_OUT)
        print(f"✅ PDF generated:  {PDF_OUT}")
    except Exception as e:
        print(f"ℹ PDF export skipped ({e}). You can export to PDF manually from Word.")

if __name__ == "__main__":
    build_doc()
