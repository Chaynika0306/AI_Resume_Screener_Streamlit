# resume_parser.py

from pathlib import Path
from pdfminer.high_level import extract_text as pdf_extract_text
import docx
import re


def extract_text_from_pdf(path):
    try:
        text = pdf_extract_text(path)
        return text or ""
    except Exception as e:
        print(f"[pdf parse error] {path}: {e}")
        return ""


def extract_text_from_docx(path):
    try:
        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs]
        return "\n".join(paragraphs)
    except Exception as e:
        print(f"[docx parse error] {path}: {e}")
        return ""


def extract_text_from_txt(path):
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        print(f"[txt parse error] {path}: {e}")
        return ""


def extract_text_from_file(file):
    """
    file: Streamlit UploadedFile OR file path
    returns: extracted text
    """

    # If file is uploaded via Streamlit
    if hasattr(file, "name"):
        filename = file.name.lower()

        if filename.endswith(".pdf"):
            from pdfminer.high_level import extract_text
            return extract_text(file)

        elif filename.endswith(".docx"):
            doc = docx.Document(file)
            return "\n".join(p.text for p in doc.paragraphs)

        elif filename.endswith(".txt"):
            return file.read().decode("utf-8", errors="ignore")

        else:
            return ""

    # If normal file path
    p = Path(file)
    suffix = p.suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(str(p))

    elif suffix in [".docx", ".doc"]:
        return extract_text_from_docx(str(p))

    elif suffix in [".txt", ".md"]:
        return extract_text_from_txt(str(p))

    return ""
# -----------------------------
# Contact Information Extractor
# -----------------------------

def extract_email(text):
    """
    Extract first email from resume text
    """
    email_pattern = r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}"

    emails = re.findall(email_pattern, text)

    if emails:
        return emails[0]

    return None


def extract_phone(text):
    """
    Extract phone number from resume text
    """
    phone_pattern = r"\+?\d[\d\s\-]{8,15}\d"

    phones = re.findall(phone_pattern, text)

    if phones:
        return phones[0]

    return None